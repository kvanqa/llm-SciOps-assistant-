"""
ingest.py

Loads documents from docs/ (.md, .txt, .pdf, .docx) and splits them into
chunks for retrieval, keeping each real section (heading + its content)
together as much as possible.

Two bugs this fixes, found via real testing against an operational document
that mixes prose with embedded Python/regex code:

1. Heading detection previously matched ANY line starting with '#', which
   false-positives on Python comments inside embedded code snippets
   (e.g. "#It specifically ignores antennas in maintenance...") — those
   got mistaken for real section headings, poisoning citations.
   Fix: readers now mark REAL headings explicitly with HEADING_MARKER at
   extraction time (from markdown '#' lines outside fenced code blocks,
   or from Word "Heading N" paragraph styles in .docx) — chunking only
   trusts these explicit markers, never guesses from '#' again.

2. Naive fixed-size sliding-window chunking doesn't respect section
   boundaries, so on a document with many short, tightly-packed
   procedures, a single chunk can span 3+ unrelated procedures, or
   truncate a short procedure's steps mid-command across a boundary.
   Fix: split into sections at real heading boundaries first; only
   sub-chunk a section with the sliding window if that section alone
   still exceeds chunk_size. Each chunk gets its heading prepended to
   the actual embedded text (not just kept as separate metadata), so
   short, code-heavy procedure chunks still carry natural-language
   signal for retrieval.
"""

import re
from dataclasses import dataclass
from pathlib import Path
from typing import List

DOCS_DIR = Path(__file__).resolve().parents[1] / "docs"

# Guaranteed not to collide with real content — used internally to mark a
# line as a REAL heading, so downstream chunking never has to guess.
HEADING_MARKER = "\x02HEADING\x02 "


@dataclass
class Chunk:
    text: str
    source: str        # filename
    chunk_index: int    # position within the source file
    heading: str = ""   # the section heading this chunk belongs to, if any


def _read_txt_or_md(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    lines = raw.splitlines()
    out = []
    in_fence = False
    for line in lines:
        if line.strip().startswith("```"):
            in_fence = not in_fence
            out.append(line)
            continue
        stripped = line.strip()
        if not in_fence and re.match(r"^#{1,6}\s+\S", stripped):
            heading_text = stripped.lstrip("#").strip()
            out.append(HEADING_MARKER + heading_text)
        else:
            out.append(line)
    return "\n".join(out)


# def _read_pdf(path: Path) -> str:
#     try:
#         from pypdf import PdfReader
#     except ImportError as e:
#         raise ImportError("Install pypdf to read PDF files: pip install pypdf") from e
#     reader = PdfReader(str(path))
#     # No reliable structural heading info from plain PDF text extraction —
#     # headings simply won't be tracked for PDFs. Worth knowing as a
#     # limitation rather than guessing at font-size heuristics.
#     return "\n\n".join(page.extract_text() or "" for page in reader.pages)

def _read_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as e:
        raise ImportError("Install pdfplumber to read PDF files with clean table structures: pip install pdfplumber") from e
    
    output_pages = []
    
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            # 1. Extract the primary text layout of the page
            page_text = page.extract_text() or ""
            
            # 2. Extract tables and format them cleanly as Markdown-style strings
            tables = page.find_tables()
            if tables:
                for table in tables:
                    raw_table_data = table.extract()
                    if raw_table_data:
                        # Convert the rows of data into a Markdown table layout
                        md_table = []
                        for i, row in enumerate(raw_table_data):
                            # Clean up None values or linebreaks inside cells
                            clean_row = [str(cell).replace('\n', ' ').strip() if cell is not None else "" for cell in row]
                            md_table.append("| " + " | ".join(clean_row) + " |")
                            
                            # Add the standard Markdown header separator line after the first row
                            if i == 0:
                                md_table.append("|" + "---| " * len(clean_row))
                        
                        # Stitch the Markdown table back into the text block
                        # You can also use a string replacement rule if you want to swap it inline,
                        # but appending it ensures the table content is intact inside the chunk.
                        page_text += "\n\n### Extracted Table Data:\n" + "\n".join(md_table)
            
            output_pages.append(page_text)
            
    return "\n\n".join(output_pages)

# def _read_pdf(path: Path) -> str:
#     '''pip uninstall torch torchaudio -y
# pip install torch torchvision torchaudio --index-url https://pytorch.org  # Change cu124 to your CUDA version
# '''
#     try:
#         from docling.document_converter import DocumentConverter
#         from docling.datamodel.accelerator_options import AcceleratorDevice, AcceleratorOptions
#         from docling.datamodel.pipeline_options import ThreadedPdfPipelineOptions
#     except ImportError as e:
#         raise ImportError("Install docling to use advanced layout extraction: pip install docling") from e
    
#     # Explicitly configure Docling to leverage your server's NVIDIA GPU via CUDA
#     acc_options = AcceleratorOptions(device=AcceleratorDevice.CUDA)
    
#     # Optimise batch sizes for GPU memory to speed up table and layout detection
#     pipeline_options = ThreadedPdfPipelineOptions(
#         ocr_batch_size=32,
#         layout_batch_size=32,
#         table_batch_size=4
#     )
    
#     # Initialize the converter with GPU settings
#     converter = DocumentConverter(
#         accelerator_options=acc_options,
#         pipeline_options=pipeline_options
#     )
    
#     # Run the AI parsing pipeline
#     result = converter.convert(path)
    
#     # Export the entire document directly as a highly structured Markdown string
#     return result.document.export_to_markdown()



def _read_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as e:
        raise ImportError("Install python-docx to read .docx files: pip install python-docx") from e
    d = docx.Document(str(path))

    parts = []
    for p in d.paragraphs:
        if not p.text.strip():
            continue
        style_name = (p.style.name or "") if p.style else ""
        if style_name.startswith("Heading") or style_name.startswith("Title"):
            parts.append(HEADING_MARKER + p.text.strip())
        else:
            parts.append(p.text)

    for table_idx, table in enumerate(d.tables):
        parts.append(f"\n{HEADING_MARKER}Table {table_idx + 1}")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


READERS = {
    ".md": _read_txt_or_md,
    ".txt": _read_txt_or_md,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}


def load_documents(docs_dir: Path = DOCS_DIR) -> List[tuple[str, str]]:
    """Returns list of (filename, full_text) for every supported file in docs_dir."""
    results = []
    for path in sorted(docs_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in READERS:
            try:
                text = READERS[path.suffix.lower()](path)
                if text.strip():
                    results.append((path.name, text))
            except Exception as e:
                print(f"[warn] skipping {path.name}: {e}")
    return results


def _split_into_sections(text: str) -> List[tuple[str, str]]:
    """
    Splits marked-up text into (heading, content) sections at HEADING_MARKER
    boundaries. Content before the first heading (if any) becomes a section
    with heading="".
    """
    lines = text.split("\n")
    sections: List[tuple[str, list[str]]] = []
    current_heading = ""
    current_lines: List[str] = []

    for line in lines:
        if line.startswith(HEADING_MARKER):
            if current_lines:
                sections.append((current_heading, current_lines))
            current_heading = line[len(HEADING_MARKER):]
            current_lines = [current_heading]  # heading text stays IN its own section's content
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_heading, current_lines))

    return [(heading, "\n".join(lines).strip()) for heading, lines in sections if "\n".join(lines).strip()]


def _sliding_window(text: str, chunk_size: int, overlap: int) -> List[str]:
    pieces = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        piece = text[start:end].strip()
        if piece:
            pieces.append(piece)
        start += chunk_size - overlap
    return pieces


def chunk_text(text: str, source: str, chunk_size: int = 800, overlap: int = 150) -> List[Chunk]:
    """
    Heading-aware chunking: split into real sections first, then only
    sub-chunk a section if it alone exceeds chunk_size. Each chunk's text
    has its heading prepended (when not already at the start), so retrieval
    has natural-language signal even for short, code-heavy procedures.
    """
    chunks: List[Chunk] = []
    idx = 0
    for heading, section_text in _split_into_sections(text):
        if len(section_text) <= chunk_size:
            pieces = [section_text]
        else:
            pieces = _sliding_window(section_text, chunk_size, overlap)

        for piece in pieces:
            embedded_text = piece if (heading and piece.startswith(heading)) else (
                f"{heading}\n\n{piece}" if heading else piece
            )
            chunks.append(Chunk(text=embedded_text, source=source, chunk_index=idx, heading=heading))
            idx += 1

    return chunks


def build_chunks(docs_dir: Path = DOCS_DIR, chunk_size: int = 800, overlap: int = 150) -> List[Chunk]:
    all_chunks: List[Chunk] = []
    for filename, text in load_documents(docs_dir):
        all_chunks.extend(chunk_text(text, filename, chunk_size, overlap))
    return all_chunks


if __name__ == "__main__":
    chunks = build_chunks()
    print(f"Loaded {len(chunks)} chunks from {DOCS_DIR}")
    for c in chunks[:3]:
        print(f"--- {c.source} [{c.heading}] chunk {c.chunk_index} ---")
        print(c.text[:200], "...\n")

